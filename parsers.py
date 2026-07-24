import re
import io
from typing import Dict, Any, List

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file using pypdf or fallback."""
    text = ""
    try:
        # pyrefly: ignore [missing-import]
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
    except Exception as e:
        # Simple regex extraction fallback for raw PDF text streams if pypdf fails
        try:
            raw_content = file_bytes.decode('utf-8', errors='ignore')
            text = " ".join(re.findall(r'[a-zA-Z0-9\s.,;:\-@()/]{4,}', raw_content))
        except Exception:
            text = ""
    return text.strip()

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file using python-docx."""
    text = ""
    try:
        # pyrefly: ignore [missing-import]
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    paragraphs.append(row_text)
        text = "\n".join(paragraphs)
    except Exception as e:
        text = file_bytes.decode('utf-8', errors='ignore')
    return text.strip()

def parse_resume_document(filename: str, file_bytes: bytes) -> Dict[str, Any]:
    """Parse resume document and extract clean text along with structured contact info."""
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        text = extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith('.docx') or filename_lower.endswith('.doc'):
        text = extract_text_from_docx(file_bytes)
    else:
        text = file_bytes.decode('utf-8', errors='ignore').strip()

    contacts = extract_contact_info(text)
    sections = extract_sections(text)
    
    return {
        "raw_text": text,
        "filename": filename,
        "contact_info": contacts,
        "sections": sections,
        "word_count": len(text.split()),
        "character_count": len(text)
    }

def extract_contact_info(text: str) -> Dict[str, str]:
    """Extract email, phone, LinkedIn, GitHub, and Portfolio URLs using regex."""
    emails = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    
    phones = re.findall(r'\(?\+?\d{1,3}\)?[\s.-]?\(?\d{2,4}\)?[\s.-]?\d{3,4}[\s.-]?\d{3,4}', text)
    # Clean phone matches
    phones = [p.strip() for p in phones if len(re.sub(r'\D', '', p)) >= 10]
    
    linkedin = re.findall(r'(?:https?://)?(?:www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+/?', text, re.IGNORECASE)
    github = re.findall(r'(?:https?://)?(?:www\.)?github\.com/[a-zA-Z0-9_-]+/?', text, re.IGNORECASE)
    
    return {
        "email": emails[0] if emails else None,
        "phone": phones[0] if phones else None,
        "linkedin": linkedin[0] if linkedin else None,
        "github": github[0] if github else None
    }

def extract_sections(text: str) -> Dict[str, str]:
    """Detect key standard resume sections (Experience, Education, Skills, Projects, Certifications)."""
    section_keywords = {
        "experience": [r'work\s+experience', r'experience', r'employment\s+history', r'work\s+history'],
        "education": [r'education', r'academic\s+background', r'qualifications'],
        "skills": [r'skills', r'technical\s+skills', r'core\s+competencies', r'technologies'],
        "projects": [r'projects', r'personal\s+projects', r'featured\s+projects'],
        "certifications": [r'certifications', r'certificates', r'licenses']
    }
    
    found_sections = {}
    lines = text.split('\n')
    current_section = "summary"
    section_buffers: Dict[str, List[str]] = {"summary": []}
    
    for line in lines:
        line_clean = line.strip().lower()
        matched = False
        
        # Check if line looks like a section header (short line matching header keywords)
        if len(line_clean) < 40:
            for sec_name, regexes in section_keywords.items():
                if any(re.search(fr'^{r}\s*$', line_clean) or re.search(fr'^{r}\s*:', line_clean) for r in regexes):
                    current_section = sec_name
                    if current_section not in section_buffers:
                        section_buffers[current_section] = []
                    matched = True
                    break
        
        if not matched:
            if current_section not in section_buffers:
                section_buffers[current_section] = []
            section_buffers[current_section].append(line)
            
    for sec, content in section_buffers.items():
        found_sections[sec] = "\n".join(content).strip()
        
    return found_sections
